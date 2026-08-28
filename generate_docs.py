"""
generate_docs.py — Generates a comprehensive, professional, dual-audience (Executives + Developers)
DOCX documentation for the TextSphere Multi-Model NLP & LLM Platform.
"""

import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls


def set_cell_background(cell, hex_color):
    """Sets background color of a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    tc_pr.append(shd)


def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    """Sets inner padding for a table cell."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = parse_xml(
        f'<w:tcMar {nsdecls("w")}>'
        f'<w:top w:w="{top}" w:type="dxa"/>'
        f'<w:bottom w:w="{bottom}" w:type="dxa"/>'
        f'<w:left w:w="{left}" w:type="dxa"/>'
        f'<w:right w:w="{right}" w:type="dxa"/>'
        f'</w:tcMar>'
    )
    tc_pr.append(tc_mar)


def set_cell_border(cell, **kwargs):
    """Sets borders for a cell (top, bottom, left, right)."""
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_borders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'<w:top w:val="{kwargs.get("top", "single")}" w:sz="{kwargs.get("top_sz", "4")}" w:space="0" w:color="{kwargs.get("top_color", "CCCCCC")}"/>'
        f'<w:left w:val="{kwargs.get("left", "none")}" w:sz="0" w:space="0" w:color="auto"/>'
        f'<w:bottom w:val="{kwargs.get("bottom", "single")}" w:sz="{kwargs.get("bottom_sz", "4")}" w:space="0" w:color="{kwargs.get("bottom_color", "CCCCCC")}"/>'
        f'<w:right w:val="{kwargs.get("right", "none")}" w:sz="0" w:space="0" w:color="auto"/>'
        f'</w:tcBorders>'
    )
    tc_pr.append(tc_borders)


def build_documentation():
    doc = Document()

    # ---------------- Page Setup (1-inch margins) ----------------
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ---------------- Style Definitions ----------------
    styles = doc.styles
    normal_style = styles['Normal']
    normal_font = normal_style.font
    normal_font.name = 'Calibri'
    normal_font.size = Pt(10.5)
    normal_font.color.rgb = RGBColor(0x22, 0x22, 0x22)
    normal_style.paragraph_format.line_spacing = 1.15
    normal_style.paragraph_format.space_after = Pt(6)

    def add_h1(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(15)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}><w:bottom w:val="single" w:sz="8" w:space="4" w:color="333333"/></w:pBdr>')
        p._p.get_or_add_pPr().append(pBdr)
        return p

    def add_h2(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(13)
        p.paragraph_format.space_after = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(12.5)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x22, 0x22, 0x22)
        return p

    def add_h3(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(9)
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(text)
        run.font.name = 'Calibri'
        run.font.size = Pt(11)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
        return p

    def add_callout(title, text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F8F9FA")
        set_cell_margins(cell, top=100, bottom=100, left=150, right=150)
        set_cell_border(cell, top="none", bottom="none", left="single", left_sz="18", left_color="222222", right="none")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        r_t = p.add_run(f"{title}: ")
        r_t.font.bold = True
        r_t.font.size = Pt(10)
        r_m = p.add_run(text)
        r_m.font.size = Pt(10)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    def add_code_block(code_text):
        tbl = doc.add_table(rows=1, cols=1)
        tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
        cell = tbl.cell(0, 0)
        set_cell_background(cell, "F4F5F7")
        set_cell_margins(cell, top=100, bottom=100, left=180, right=180)
        set_cell_border(cell, top="single", top_sz="4", top_color="D0D5DD",
                              bottom="single", bottom_sz="4", bottom_color="D0D5DD")
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.0
        run = p.add_run(code_text.strip())
        run.font.name = 'Consolas'
        run.font.size = Pt(8.5)
        run.font.color.rgb = RGBColor(0x1F, 0x24, 0x2E)
        doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # ---------------- DOCUMENT TITLE & HEADER ----------------
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(20)
    title_p.paragraph_format.space_after = Pt(4)
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    t_run = title_p.add_run("TEXTSPHERE")
    t_run.font.name = 'Calibri'
    t_run.font.size = Pt(26)
    t_run.font.bold = True
    t_run.font.color.rgb = RGBColor(0x11, 0x11, 0x11)

    sub_p = doc.add_paragraph()
    sub_p.paragraph_format.space_before = Pt(0)
    sub_p.paragraph_format.space_after = Pt(14)
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run("Product Overview, System Architecture & Technical Developer Manual")
    sub_run.font.name = 'Calibri'
    sub_run.font.size = Pt(13)
    sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    # Metadata Table
    meta_table = doc.add_table(rows=5, cols=2)
    meta_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta_data = [
        ("Project Name", "TextSphere NLP & LLM Platform"),
        ("Audience", "Project Evaluators, Business Stakeholders, ML Engineers, Full-Stack Developers"),
        ("System Architecture", "Dual-Engine (Local Fine-Tuned Transformer Models + Cloud LLM Streaming)"),
        ("Technology Stack", "FastAPI, PyTorch, HuggingFace Transformers, React 18, Vite, Firebase"),
        ("Repository", "https://github.com/AbdullahRandhawa/TextSphere"),
    ]
    for row_idx, (k, v) in enumerate(meta_data):
        cell_k = meta_table.cell(row_idx, 0)
        cell_v = meta_table.cell(row_idx, 1)
        cell_k.width = Inches(2.2)
        cell_v.width = Inches(4.3)
        set_cell_background(cell_k, "F8F9FA")
        set_cell_margins(cell_k, top=50, bottom=50, left=90, right=90)
        set_cell_margins(cell_v, top=50, bottom=50, left=90, right=90)
        set_cell_border(cell_k, top="single", bottom="single", top_color="E0E0E0", bottom_color="E0E0E0")
        set_cell_border(cell_v, top="single", bottom="single", top_color="E0E0E0", bottom_color="E0E0E0")

        pk = cell_k.paragraphs[0]
        pk.paragraph_format.space_after = Pt(0)
        rk = pk.add_run(k)
        rk.font.bold = True
        rk.font.size = Pt(9.0)

        pv = cell_v.paragraphs[0]
        pv.paragraph_format.space_after = Pt(0)
        rv = pv.add_run(v)
        rv.font.size = Pt(9.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(10)

    # =========================================================================
    # PART I: GENERAL & EXECUTIVE PRODUCT OVERVIEW
    # =========================================================================
    add_h1("PART I: PRODUCT & BUSINESS OVERVIEW")

    add_h2("1. What is TextSphere?")
    doc.add_paragraph(
        "TextSphere is an all-in-one Natural Language Processing (NLP) and Artificial Intelligence workspace. "
        "It bridges the gap between everyday conversational AI (like ChatGPT) and high-precision, domain-specific "
        "machine learning tools."
    )
    doc.add_paragraph(
        "In traditional systems, companies either rely exclusively on massive cloud LLMs (which are slow, expensive, "
        "and prone to hallucinating factual labels) or use isolated ML scripts without any conversational interface. "
        "TextSphere introduces a unified Dual-Engine Architecture where users can chat normally, but also trigger "
        "specialized, fine-tuned transformer models on-demand. The local models calculate exact mathematical predictions "
        "and confidence scores, while the LLM instantly analyzes and explains the findings in real-time streaming language."
    )

    add_h2("2. Key Problems Solved")
    prob_items = [
        ("High Cloud API Costs", "Routing routine tasks like sentiment detection or text classification to commercial LLMs incurs continuous per-token costs. TextSphere runs these on local fine-tuned transformers at zero marginal token cost."),
        ("Slow Latency & Network Bottlenecks", "Local transformer inference executes in milliseconds directly in PyTorch memory, avoiding external API round-trip delays."),
        ("Lack of Mathematical Confidence", "General LLMs generate text without verifiable certainty scores. TextSphere's local models output calibrated probability scores (e.g., 99.8% positive, 98.5% Sci/Tech) for reliable decision-making."),
        ("Scattered Tooling", "Instead of using five separate utilities for summarization, entity extraction, sentiment, question answering, and topic labeling, TextSphere consolidates everything into a single modern web workspace."),
    ]
    for title, desc in prob_items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f"{title}: ")
        r.font.bold = True
        p.add_run(desc)

    add_h2("3. The 5 Core NLP Tools Explained (For Non-Technical Users)")
    doc.add_paragraph(
        "TextSphere is equipped with five specialized AI engines built right into the interface:"
    )

    tools_overview = [
        ("1. Sentiment Analyzer", "Determines whether an input text expresses a Positive or Negative tone along with a numerical confidence percentage. Perfect for customer review analysis and social media monitoring."),
        ("2. Topic Classifier", "Automatically categorizes articles and documents into four primary news/industry domains: World News, Sports, Business, and Science/Technology."),
        ("3. Named Entity Recognizer (NER)", "Scans text and identifies real-world entities such as People (PER), Organizations (ORG), and Locations (LOC), highlighting exact character positions."),
        ("4. Text Summarizer", "Condenses lengthy articles, transcripts, or reports into clear, concise summaries while retaining critical factual context."),
        ("5. Question Answering Engine", "Extracts direct, pinpoint answers to user questions from a provided reference text or background document without hallucinating outside information."),
    ]
    for title, desc in tools_overview:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f"{title}: ")
        r.font.bold = True
        p.add_run(desc)

    # =========================================================================
    # PART II: MACHINE LEARNING & MODEL ENGINEERING
    # =========================================================================
    add_h1("PART II: MACHINE LEARNING & MODEL ENGINEERING")

    doc.add_paragraph(
        "Each tool in TextSphere is backed by a dedicated transformer neural network fine-tuned on "
        "gold-standard research datasets. All model checkpoints are stored and executed using Safetensors weights."
    )

    # Table of models
    tbl_models = doc.add_table(rows=6, cols=5)
    tbl_models.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ["Tool Name", "Base Architecture", "Training Dataset", "Model Output", "Sample Live Output"]
    for c_idx, h in enumerate(headers):
        cell = tbl_models.cell(0, c_idx)
        set_cell_background(cell, "1F242E")
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    models_data = [
        ("Sentiment Analyzer", "DistilBERT (uncased)", "SST-2 (Stanford Sentiment)", "Polarity + Prob.", "{'label': 'Positive', 'confidence': 0.9987}"),
        ("Topic Classifier", "DistilBERT (uncased)", "AG News (120k news items)", "Class + Prob.", "{'label': 'Sci/Tech', 'confidence': 0.9853}"),
        ("Named Entity (NER)", "BERT (cased)", "CoNLL-2003 (Reuters)", "Entity Spans", "[{'text': 'Bill Gates', 'label': 'PER'}]"),
        ("Summarizer", "T5-small (Seq2Seq)", "CNN / DailyMail", "Abstractive Text", "{'summary': 'FastAPI is a fast web framework...'}"),
        ("Question Answering", "DistilBERT (cased)", "SQuAD v1.1 (100k+ Q&As)", "Extractive Span", "{'answer': 'Microsoft', 'confidence': 0.9976}"),
    ]

    for r_idx, row_data in enumerate(models_data, start=1):
        bg = "FFFFFF" if r_idx % 2 != 0 else "F8F9FA"
        for c_idx, val in enumerate(row_data):
            cell = tbl_models.cell(r_idx, c_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=70, right=70)
            set_cell_border(cell, top="single", bottom="single", top_color="E0E0E0", bottom_color="E0E0E0")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(8.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_h2("4. In-Place Memory Loader & Zero-Copy Architecture")
    doc.add_paragraph(
        "To maximize efficiency and eliminate unnecessary disk usage, the system features an optimized "
        "in-place loader (backend/app/tools/_loader.py). Rather than copying weights into temporary OS directories, "
        "the loader initializes weights directly from backend/app/finetuned_models/ in-place. "
        "Models are cached for the process lifetime, ensuring instantaneous subsequent inference."
    )

    add_h2("5. Model Storage & Google Drive Synchronization")
    doc.add_paragraph(
        "Because transformer checkpoints total ~1.4 GB (exceeding standard Git limits), they are securely hosted "
        "on Google Drive and tracked via an automated downloader (download_models.py):"
    )
    p_d1 = doc.add_paragraph(style='List Bullet')
    p_d1.add_run("Sentiment Model: Drive ID 10R9YmDnIKz9XgCiCUqXszl7l4JS_6Y47")
    p_d2 = doc.add_paragraph(style='List Bullet')
    p_d2.add_run("Topic Model: Drive ID 15lz2jiavSmRKvzZF5j_wyaTFYdne5cjZ")
    p_d3 = doc.add_paragraph(style='List Bullet')
    p_d3.add_run("NER Model: Drive ID 1RZVF5SEdh6p3wtch9Ep9dzCgh5Det2KA")
    p_d4 = doc.add_paragraph(style='List Bullet')
    p_d4.add_run("Summarization Model: Drive ID 1F8wuov-ro9yx-5p7qZvh7OLky-xSZzjP")
    p_d5 = doc.add_paragraph(style='List Bullet')
    p_d5.add_run("QA Model: Drive ID 15nZvciSd6tNZ4QoVJFkw34M4C62CbSIV")

    # =========================================================================
    # PART III: SOFTWARE ARCHITECTURE & DEVELOPER GUIDE
    # =========================================================================
    add_h1("PART III: SOFTWARE ARCHITECTURE & DEVELOPER GUIDE")

    add_h2("6. System Architecture & Information Flow")
    doc.add_paragraph(
        "TextSphere follows a modern decoupled architecture: React 18 SPA on the frontend, FastAPI on the backend, "
        "and Google Firebase for auth & database services."
    )

    add_code_block("""[User Browser: React 18 + Vite]
        │
        ├── 1. Authenticates via Firebase Client SDK (OAuth/Password)
        │
        └── 2. Emits POST /chat (with Firebase JWT + Message + Optional toolId)
                │
                ▼
        [Backend: FastAPI Async Engine]
                ├── A. Verifies JWT Bearer Token (Firebase Admin SDK)
                ├── B. Executes Rate Limit check in Firestore
                ├── C. Runs Local PyTorch Model (Zero Token Cost) -> Emits 'tool_result' SSE
                ├── D. Calls OpenRouter LLM Stream -> Emits 'commentary_chunk' SSE
                ├── E. Concurrently writes message history to Firestore
                └── F. Emits 'done' SSE event""")

    add_h2("7. Server-Sent Events (SSE) Protocol")
    doc.add_paragraph(
        "The POST /chat endpoint streams real-time data using the SSE text/event-stream specification. "
        "Events are structured as JSON strings prefixed by 'data: ':"
    )

    add_code_block("""// Event 1: Local Tool Prediction Output
data: {
  "event": "tool_result",
  "tool_id": "sentiment",
  "display_name": "Sentiment Analyzer",
  "base_model": "DistilBERT",
  "fine_tune_dataset": "SST-2",
  "result": {"label": "Positive", "confidence": 0.9987}
}

// Event 2: Streaming LLM Analysis
data: {"event": "commentary_chunk", "text": "The analysis indicates strong positive sentiment..."}

// Event 3: Stream Finalized
data: {"event": "done"}""")

    add_h2("8. Database Schema & Persistence (Cloud Firestore)")
    doc.add_paragraph(
        "Data isolation is maintained by scoping all collections under authenticated user IDs:"
    )

    tbl_fs = doc.add_table(rows=3, cols=3)
    tbl_fs.alignment = WD_TABLE_ALIGNMENT.CENTER
    fs_headers = ["Firestore Path", "Stored Fields", "Function"]
    for c_idx, h in enumerate(fs_headers):
        cell = tbl_fs.cell(0, c_idx)
        set_cell_background(cell, "1F242E")
        set_cell_margins(cell, top=90, bottom=90, left=90, right=90)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        r = p.add_run(h)
        r.font.bold = True
        r.font.size = Pt(8.5)
        r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    fs_data = [
        ("users/{uid}/chats/{chatId}", "title, createdAt, updatedAt, messageCount, apiCallCount", "Tracks chat sessions, conversation metadata, and rate limit counters."),
        ("users/{uid}/chats/{chatId}/messages/{msgId}", "role ('user'|'assistant'), content, toolUsed, toolResult, timestamp", "Individual message history. Stores full structured tool payloads alongside LLM commentary."),
    ]
    for r_idx, (path, fields, desc) in enumerate(fs_data, start=1):
        bg = "FFFFFF" if r_idx % 2 != 0 else "F8F9FA"
        for c_idx, val in enumerate([path, fields, desc]):
            cell = tbl_fs.cell(r_idx, c_idx)
            set_cell_background(cell, bg)
            set_cell_margins(cell, top=70, bottom=70, left=70, right=70)
            set_cell_border(cell, top="single", bottom="single", top_color="E0E0E0", bottom_color="E0E0E0")
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            r = p.add_run(val)
            r.font.size = Pt(8.0)

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_h2("9. Frontend Architecture (React 18 Custom Hooks)")
    hooks_list = [
        ("useStreamingChat.js", "Coordinates the SSE stream parser, handles chunk accumulation, auto-scrolling, and manages loading indicators."),
        ("useChats.js", "Manages chat state, real-time Firestore listeners, session creation, switching, and deletion."),
        ("useAuth.js", "Wraps Firebase Authentication, providing reactive user states, login/logout functions, and token management."),
    ]
    for name, desc in hooks_list:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(f"{name}: ")
        r.font.bold = True
        p.add_run(desc)

    add_h2("10. Installation & Quick Start Guide")
    doc.add_paragraph("Follow these steps to launch the system on any development environment:")
    
    add_code_block("""# 1. Clone repository
git clone https://github.com/AbdullahRandhawa/TextSphere.git
cd TextSphere

# 2. Download fine-tuned model checkpoints
pip install gdown
python download_models.py

# 3. Verify environment & dependencies
python setup_check.py

# 4. Start Backend Server (Port 8000)
cd backend && uvicorn app.main:app --reload

# 5. Start Frontend UI (Port 5173)
cd frontend && npm install && npm run dev""")

    # Save document (handles file lock if Word has it open)
    output_path = "TextSphere_Documentation.docx"
    doc.save(output_path)
    print(f"Successfully generated: {output_path}")
    
    try:
        doc.save("TextSphere_Developer_Documentation.docx")
        print("Also updated: TextSphere_Developer_Documentation.docx")
    except Exception:
        print("Note: TextSphere_Developer_Documentation.docx is currently open in Word. Saved as TextSphere_Documentation.docx instead.")


if __name__ == "__main__":
    build_documentation()
